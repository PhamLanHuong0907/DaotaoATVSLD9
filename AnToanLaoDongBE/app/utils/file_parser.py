import os
import math
import logging

import pdfplumber
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


def _rows_to_markdown_table(rows: list[list]) -> str:
    """Convert a 2D list of cell values into a GitHub-flavored Markdown table.

    Empty / None cells → blank. Trims whitespace and escapes pipes.
    Returns '' if fewer than 2 rows (not a useful table).
    """
    clean_rows: list[list[str]] = []
    for r in rows:
        cells = []
        for c in r:
            s = "" if c is None else str(c).replace("|", "\\|").replace("\n", " ").strip()
            cells.append(s)
        clean_rows.append(cells)

    # Normalize column count
    max_cols = max((len(r) for r in clean_rows), default=0)
    if max_cols < 2 or len(clean_rows) < 2:
        return ""

    for r in clean_rows:
        while len(r) < max_cols:
            r.append("")

    header = clean_rows[0]
    sep = ["---"] * max_cols
    body = clean_rows[1:]

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ]
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


class ParseResult:
    """Result of file parsing with page-level granularity."""

    def __init__(self, pages: list[str], page_count: int | None = None):
        self.pages = pages
        self.page_count = page_count or len(pages)

    @property
    def full_text(self) -> str:
        return "\n\n".join(self.pages)

    @property
    def total_chars(self) -> int:
        return sum(len(p) for p in self.pages)

    def get_chunks(self, max_chars_per_chunk: int = 30000) -> list[str]:
        """
        Split pages into chunks, each under max_chars_per_chunk.
        Keeps whole pages together - never splits mid-page.
        """
        chunks = []
        current_chunk_parts = []
        current_size = 0

        for page in self.pages:
            page_len = len(page)

            # If a single page exceeds the limit, it becomes its own chunk
            if page_len > max_chars_per_chunk:
                # Flush current chunk first
                if current_chunk_parts:
                    chunks.append("\n\n".join(current_chunk_parts))
                    current_chunk_parts = []
                    current_size = 0
                # Split the large page into sub-chunks by paragraphs
                for sub in _split_large_text(page, max_chars_per_chunk):
                    chunks.append(sub)
                continue

            # Adding this page would exceed limit -> flush current chunk
            if current_size + page_len > max_chars_per_chunk and current_chunk_parts:
                chunks.append("\n\n".join(current_chunk_parts))
                current_chunk_parts = []
                current_size = 0

            current_chunk_parts.append(page)
            current_size += page_len

        # Flush remaining
        if current_chunk_parts:
            chunks.append("\n\n".join(current_chunk_parts))

        return chunks


def _split_large_text(text: str, max_chars: int) -> list[str]:
    """Split a very large text block by paragraphs, keeping under max_chars."""
    paragraphs = text.split("\n")
    chunks = []
    current = []
    current_size = 0

    for para in paragraphs:
        if current_size + len(para) > max_chars and current:
            chunks.append("\n".join(current))
            current = []
            current_size = 0
        current.append(para)
        current_size += len(para)

    if current:
        chunks.append("\n".join(current))
    return chunks


class FileParser:
    MIME_PDF = "application/pdf"
    MIME_DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    MIME_DOC = "application/msword"
    MIME_XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    MIME_XLS = "application/vnd.ms-excel"
    MIME_TXT = "text/plain"

    @staticmethod
    def extract(file_path: str, mime_type: str) -> ParseResult:
        """Extract text from file, returning page-level results."""
        if mime_type == FileParser.MIME_PDF:
            return FileParser._extract_pdf(file_path)
        elif mime_type in (FileParser.MIME_DOCX, FileParser.MIME_DOC):
            return FileParser._extract_docx(file_path)
        elif mime_type in (FileParser.MIME_XLSX, FileParser.MIME_XLS):
            return FileParser._extract_xlsx(file_path)
        elif mime_type.startswith("text/"):
            return FileParser._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    # Keep backward-compatible method
    @staticmethod
    def extract_text(file_path: str, mime_type: str) -> tuple[str, int | None]:
        result = FileParser.extract(file_path, mime_type)
        return result.full_text, result.page_count

    @staticmethod
    def _extract_pdf(file_path: str) -> ParseResult:
        """Trích văn bản + bảng (dưới dạng Markdown) từ PDF bằng pdfplumber.

        Chiến lược:
          - Dò bảng với pdfplumber.find_tables() → convert sang Markdown.
          - Lấy text còn lại của trang, loại bỏ vùng đã dùng cho bảng.
          - Fallback về PyPDF2 nếu pdfplumber lỗi.
        """
        pages: list[str] = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    parts: list[str] = []

                    # 1) Bảng → Markdown (chèn vào đúng vị trí top-down)
                    tables = []
                    try:
                        raw_tables = page.find_tables()
                        for t in raw_tables:
                            rows = t.extract()
                            md = _rows_to_markdown_table(rows or [])
                            if md:
                                tables.append({"top": t.bbox[1], "md": md, "bbox": t.bbox})
                    except Exception as e:
                        logger.debug(f"PDF table detection failed p{i+1}: {e}")
                        tables = []

                    # 2) Text thô (loại vùng có bảng)
                    try:
                        if tables:
                            # Loại vùng có bảng khỏi text chính
                            bboxes = [t["bbox"] for t in tables]
                            non_table = page
                            for bbox in bboxes:
                                try:
                                    non_table = non_table.outside_bbox(bbox)
                                except Exception:
                                    pass
                            page_text = non_table.extract_text() or ""
                        else:
                            page_text = page.extract_text() or ""
                    except Exception as e:
                        logger.debug(f"pdfplumber extract_text failed p{i+1}: {e}")
                        page_text = ""

                    page_text = page_text.strip()

                    # 3) Ghép text + bảng (bảng ở cuối trang để không vỡ dòng)
                    content = ""
                    if page_text:
                        content = page_text
                    if tables:
                        tables.sort(key=lambda x: x["top"])
                        table_blocks = "\n\n".join(t["md"] for t in tables)
                        content = f"{content}\n\n{table_blocks}" if content else table_blocks

                    if content.strip():
                        pages.append(f"[Trang {i + 1}]\n{content}")

                page_count = len(pdf.pages)
        except Exception as e:
            logger.warning(f"pdfplumber failed, fallback PyPDF2: {e}")
            reader = PdfReader(file_path)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append(f"[Trang {i + 1}]\n{text}")
            page_count = len(reader.pages)

        return ParseResult(pages=pages, page_count=page_count)

    @staticmethod
    def _extract_docx(file_path: str) -> ParseResult:
        """Trích đoạn văn + bảng (Markdown) từ DOCX theo đúng thứ tự xuất hiện."""
        doc = DocxDocument(file_path)

        # Duyệt body theo thứ tự XML để giữ đúng xen kẽ paragraph/table
        from docx.oxml.ns import qn
        body = doc.element.body
        para_map = {p._p: p for p in doc.paragraphs}
        table_map = {t._tbl: t for t in doc.tables}

        blocks: list[str] = []
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                p = para_map.get(child)
                if p is None:
                    continue
                text = p.text.strip()
                if text:
                    blocks.append(text)
            elif child.tag == qn("w:tbl"):
                tbl = table_map.get(child)
                if tbl is None:
                    continue
                rows = [[cell.text for cell in row.cells] for row in tbl.rows]
                md = _rows_to_markdown_table(rows)
                if md:
                    blocks.append(md)

        # Nhóm khoảng 2000 ký tự thành 1 "trang"
        pages = []
        current_page: list[str] = []
        current_size = 0
        page_num = 1
        for blk in blocks:
            current_page.append(blk)
            current_size += len(blk)
            if current_size >= 2000:
                pages.append(f"[Phần {page_num}]\n" + "\n\n".join(current_page))
                current_page = []
                current_size = 0
                page_num += 1
        if current_page:
            pages.append(f"[Phần {page_num}]\n" + "\n\n".join(current_page))

        return ParseResult(pages=pages)

    @staticmethod
    def _extract_xlsx(file_path: str) -> ParseResult:
        wb = load_workbook(file_path, read_only=True)
        pages = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = [list(row) for row in ws.iter_rows(values_only=True) if any(c is not None for c in row)]
            md = _rows_to_markdown_table(rows)
            if md:
                pages.append(f"[Sheet: {sheet_name}]\n{md}")
        wb.close()
        return ParseResult(pages=pages)

    @staticmethod
    def _extract_txt(file_path: str) -> ParseResult:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        # Split into ~2000 char chunks
        chunk_size = 2000
        pages = []
        for i in range(0, len(text), chunk_size):
            chunk = text[i:i + chunk_size]
            if chunk.strip():
                pages.append(chunk)
        return ParseResult(pages=pages if pages else [text])

    @staticmethod
    def get_mime_type(filename: str) -> str:
        ext = os.path.splitext(filename)[1].lower()
        mime_map = {
            ".pdf": FileParser.MIME_PDF,
            ".docx": FileParser.MIME_DOCX,
            ".doc": FileParser.MIME_DOC,
            ".xlsx": FileParser.MIME_XLSX,
            ".xls": FileParser.MIME_XLS,
            ".txt": FileParser.MIME_TXT,
        }
        return mime_map.get(ext, "application/octet-stream")
